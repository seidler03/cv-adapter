"""
Generate a downloadable DOCX from adapted CV text.
"""
import io


def build_docx(cv_text: str, title: str = 'Adapted CV') -> bytes:
    """Return bytes of a DOCX file from plain text."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

        for line in cv_text.split('\n'):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph()
                continue
            # Heuristic: all-caps lines or lines ending with ':' → heading
            if stripped.isupper() or (len(stripped) < 60 and stripped.endswith(':')):
                p = doc.add_heading(stripped, level=2)
            else:
                p = doc.add_paragraph(stripped)
                p.style.font.size = Pt(11)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    except ImportError:
        raise ImportError("python-docx is required. Run: pip install python-docx")
