"""
CV text extraction from PDF and DOCX files.
"""
import io


def extract_text_from_pdf(file_obj) -> str:
    """Extract text from a PDF using PyMuPDF (primary) with pdfplumber fallback."""
    data = file_obj.read() if hasattr(file_obj, 'read') else open(file_obj, 'rb').read()

    # Primary: PyMuPDF (fast, handles most PDFs)
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=data, filetype='pdf')
        pages = [page.get_text() for page in doc]
        doc.close()
        text = '\n'.join(pages).strip()
        if text:
            return text
    except ImportError:
        pass  # fall through to pdfplumber
    except Exception:
        pass  # corrupted / encrypted page — try fallback

    # Fallback: pdfplumber (better for columnar/table layouts)
    try:
        import pdfplumber
        import io
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            pages = [page.extract_text() or '' for page in pdf.pages]
        text = '\n'.join(pages).strip()
        if text:
            return text
    except ImportError:
        pass
    except Exception as exc:
        raise ValueError(f"Failed to parse PDF: {exc}") from exc

    raise ValueError(
        "Could not extract text from this PDF. "
        "It may be scanned/image-only. Please upload a text-based PDF or a DOCX file."
    )


def extract_text_from_docx(file_obj) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
        data = file_obj.read() if hasattr(file_obj, 'read') else None
        if data:
            doc = Document(io.BytesIO(data))
        else:
            doc = Document(file_obj)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return '\n'.join(paragraphs).strip()
    except ImportError:
        raise ImportError(
            "python-docx is not installed. Run: pip install python-docx"
        )
    except Exception as exc:
        raise ValueError(f"Failed to parse DOCX: {exc}") from exc


def extract_text(file_obj, filename: str) -> str:
    """Auto-detect file type and extract text."""
    ext = filename.rsplit('.', 1)[-1].lower()
    if ext == 'pdf':
        return extract_text_from_pdf(file_obj)
    elif ext in ('docx', 'doc'):
        return extract_text_from_docx(file_obj)
    else:
        raise ValueError(f"Unsupported file type: .{ext}")
